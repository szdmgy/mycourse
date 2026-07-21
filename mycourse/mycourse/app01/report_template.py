# -*- coding: utf-8 -*-
"""作业报告模板：路径、存取与封面自动填写（保留 Word run/下划线格式）。"""
from __future__ import annotations

import io
import os
import re
import shutil
from dataclasses import dataclass

from django.conf import settings
from django.utils import timezone
from docx import Document

from app01.utils import safe_filename

TEMPLATE_MAX_BYTES = 30 * 1024 * 1024
TEMPLATE_MAX_MB = TEMPLATE_MAX_BYTES // (1024 * 1024)


# ──────────────────────────── 路径 ────────────────────────────

def report_template_rel_dir(course, task_title: str) -> str:
    return os.path.join(
        'file',
        course.courseTerm,
        course.courseName + course.classNumber,
        '报告模板',
        safe_filename(task_title),
    )


def course_default_template_rel_dir(course) -> str:
    return os.path.join(
        'file',
        course.courseTerm,
        course.courseName + course.classNumber,
        '报告模板',
        '_课程默认',
    )


def template_abs_path(rel_or_abs: str) -> str:
    if os.path.isabs(rel_or_abs):
        return rel_or_abs
    return os.path.join(settings.BASE_DIR, rel_or_abs)


def _path_has_file(rel: str) -> bool:
    rel = (rel or '').strip()
    return bool(rel) and os.path.isfile(template_abs_path(rel))


def task_has_template_file(task) -> bool:
    return _path_has_file(getattr(task, 'template_path', ''))


def course_has_template_file(course) -> bool:
    return _path_has_file(getattr(course, 'report_template_path', ''))


@dataclass
class ResolvedTemplate:
    abs_path: str
    original_name: str
    enable_download: bool
    enable_autofill: bool
    source: str  # 'task' | 'course'


def resolve_task_template(task) -> ResolvedTemplate | None:
    """作业专用模板优先；否则回落到课程默认模板。纯非 docx 作业不解析模板。"""
    if not task_allows_report_template(task):
        return None
    if task_has_template_file(task):
        return ResolvedTemplate(
            abs_path=template_abs_path(task.template_path),
            original_name=task.template_original_name or os.path.basename(task.template_path),
            enable_download=bool(task.enable_template_download),
            enable_autofill=bool(task.enable_cover_autofill),
            source='task',
        )
    course = task.courseBelongTo
    if course_has_template_file(course):
        return ResolvedTemplate(
            abs_path=template_abs_path(course.report_template_path),
            original_name=course.report_template_original_name or os.path.basename(course.report_template_path),
            enable_download=bool(course.enable_report_template_download),
            enable_autofill=bool(course.enable_report_cover_autofill),
            source='course',
        )
    return None


def template_available_for_student(task) -> bool:
    if not task_allows_report_template(task):
        return False
    resolved = resolve_task_template(task)
    return bool(resolved and resolved.enable_download)


def effective_cover_autofill(task) -> bool:
    if not task_allows_report_template(task):
        return False
    resolved = resolve_task_template(task)
    return bool(resolved and resolved.enable_download and resolved.enable_autofill)


def _remove_file_quiet(abs_path: str) -> None:
    try:
        if os.path.isfile(abs_path):
            os.remove(abs_path)
    except OSError:
        pass
    parent = os.path.dirname(abs_path)
    try:
        if os.path.isdir(parent) and not os.listdir(parent):
            os.rmdir(parent)
    except OSError:
        pass


def clear_task_template_files(task) -> None:
    path = (getattr(task, 'template_path', None) or '').strip()
    if path:
        _remove_file_quiet(template_abs_path(path))


def clear_course_template_files(course) -> None:
    path = (getattr(course, 'report_template_path', None) or '').strip()
    if path:
        _remove_file_quiet(template_abs_path(path))


def _write_upload(abs_path: str, uploaded_file) -> None:
    os.makedirs(os.path.dirname(abs_path), exist_ok=True)
    with open(abs_path, 'wb') as dest:
        for chunk in uploaded_file.chunks():
            dest.write(chunk)


def _validate_docx_upload(uploaded_file) -> tuple[str | None, str | None]:
    name = uploaded_file.name or 'template.docx'
    base = os.path.basename(name)
    if base.lower().endswith('.doc') and not base.lower().endswith('.docx'):
        return '不支持旧版 Word（.doc），模板请使用 .docx 格式', None
    if not base.lower().endswith('.docx'):
        return '仅支持 .docx 模板文件', None
    size = getattr(uploaded_file, 'size', None)
    if size is not None and size > TEMPLATE_MAX_BYTES:
        return f'模板文件不能超过 {TEMPLATE_MAX_MB}MB', None
    return None, base


def save_uploaded_template(task, uploaded_file) -> tuple[str | None, str | None]:
    if not task_allows_report_template(task):
        return '本作业不允许提交 .docx（如仅 zip），不能挂报告模板', None
    err, base = _validate_docx_upload(uploaded_file)
    if err:
        return err, None
    clear_task_template_files(task)
    rel_dir = report_template_rel_dir(task.courseBelongTo, task.title)
    safe_base = safe_filename(base)
    if not safe_base.lower().endswith('.docx'):
        safe_base += '.docx'
    rel_path = os.path.join(rel_dir, safe_base)
    _write_upload(os.path.join(settings.BASE_DIR, rel_path), uploaded_file)
    task.template_path = rel_path
    task.template_original_name = base
    task.template_uploaded_at = timezone.now()
    task.enable_template_download = True
    task.enable_cover_autofill = True
    task.save(update_fields=[
        'template_path', 'template_original_name', 'template_uploaded_at',
        'enable_template_download', 'enable_cover_autofill',
    ])
    return None, base


def save_uploaded_course_template(course, uploaded_file) -> tuple[str | None, str | None]:
    err, base = _validate_docx_upload(uploaded_file)
    if err:
        return err, None
    clear_course_template_files(course)
    rel_dir = course_default_template_rel_dir(course)
    safe_base = safe_filename(base)
    if not safe_base.lower().endswith('.docx'):
        safe_base += '.docx'
    rel_path = os.path.join(rel_dir, safe_base)
    _write_upload(os.path.join(settings.BASE_DIR, rel_path), uploaded_file)
    course.report_template_path = rel_path
    course.report_template_original_name = base
    course.report_template_uploaded_at = timezone.now()
    course.enable_report_template_download = True
    course.enable_report_cover_autofill = True
    course.save(update_fields=[
        'report_template_path', 'report_template_original_name', 'report_template_uploaded_at',
        'enable_report_template_download', 'enable_report_cover_autofill',
    ])
    return None, base


def copy_template_to_task(src_task, dest_task) -> bool:
    if not task_has_template_file(src_task):
        dest_task.template_path = ''
        dest_task.template_original_name = ''
        dest_task.template_uploaded_at = None
        dest_task.enable_template_download = False
        dest_task.enable_cover_autofill = False
        dest_task.save(update_fields=[
            'template_path', 'template_original_name', 'template_uploaded_at',
            'enable_template_download', 'enable_cover_autofill',
        ])
        return False

    src_abs = template_abs_path(src_task.template_path)
    orig = src_task.template_original_name or os.path.basename(src_task.template_path)
    safe_base = safe_filename(orig)
    if not safe_base.lower().endswith('.docx'):
        safe_base += '.docx'
    rel_dir = report_template_rel_dir(dest_task.courseBelongTo, dest_task.title)
    abs_dir = os.path.join(settings.BASE_DIR, rel_dir)
    os.makedirs(abs_dir, exist_ok=True)
    rel_path = os.path.join(rel_dir, safe_base)
    abs_path = os.path.join(settings.BASE_DIR, rel_path)
    shutil.copy2(src_abs, abs_path)
    dest_task.template_path = rel_path
    dest_task.template_original_name = orig
    dest_task.template_uploaded_at = src_task.template_uploaded_at
    dest_task.enable_template_download = src_task.enable_template_download
    dest_task.enable_cover_autofill = src_task.enable_cover_autofill
    dest_task.save(update_fields=[
        'template_path', 'template_original_name', 'template_uploaded_at',
        'enable_template_download', 'enable_cover_autofill',
    ])
    return True


def copy_course_template_to_course(src_course, dest_course) -> bool:
    """历史开班时可选：复制课程默认模板。"""
    if not course_has_template_file(src_course):
        return False
    src_abs = template_abs_path(src_course.report_template_path)
    orig = src_course.report_template_original_name or os.path.basename(src_course.report_template_path)
    safe_base = safe_filename(orig)
    if not safe_base.lower().endswith('.docx'):
        safe_base += '.docx'
    clear_course_template_files(dest_course)
    rel_dir = course_default_template_rel_dir(dest_course)
    rel_path = os.path.join(rel_dir, safe_base)
    abs_path = os.path.join(settings.BASE_DIR, rel_path)
    os.makedirs(os.path.dirname(abs_path), exist_ok=True)
    shutil.copy2(src_abs, abs_path)
    dest_course.report_template_path = rel_path
    dest_course.report_template_original_name = orig
    dest_course.report_template_uploaded_at = src_course.report_template_uploaded_at
    dest_course.enable_report_template_download = src_course.enable_report_template_download
    dest_course.enable_report_cover_autofill = src_course.enable_report_cover_autofill
    dest_course.save(update_fields=[
        'report_template_path', 'report_template_original_name', 'report_template_uploaded_at',
        'enable_report_template_download', 'enable_report_cover_autofill',
    ])
    return True


# ──────────────────────────── 封面填写（保留格式） ────────────────────────────

# 深圳大学实验报告封面：学院/专业固定文案（标准模板）
FIXED_COLLEGE = '生物医学工程学院'
FIXED_MAJOR = '生物医学工程专业'

# 永不自动填写的封面项
SKIP_COVER_LABELS = frozenset({'实验时间', '实验报告提交时间', '提交时间'})


def task_allows_report_template(task) -> bool:
    """报告模板仅对允许提交 .docx 的作业生效；纯 zip 等作业不允许挂模板。"""
    ft = (getattr(task, 'fileType', None) or '*').strip().lower()
    if ft in ('', '*'):
        return True
    allowed = [x.strip().lstrip('.') for x in ft.split(',') if x.strip()]
    return 'docx' in allowed


def build_cover_values(task, profile, user) -> dict:
    course = task.courseBelongTo
    class_no = (course.classNumber or '').strip()
    class_label = class_no if class_no.endswith('班') else (f'{class_no}班' if class_no else '')
    name = (profile.name or '').strip()
    number = (user.username or '').strip()
    return {
        '姓名': name,
        '学号': number,
        '班级': class_label,
        '课程名称': course.courseName or '',
        '实验名称': task.title or '',
        '作业标题': task.title or '',
        '指导教师': (course.teachers or '').strip(),
        '学期': course.courseTerm or '',
        '学院': FIXED_COLLEGE,
        '专业': FIXED_MAJOR,
        'name': name,
        'student_number': number,
        'class': class_label,
        'course_name': course.courseName or '',
        'task_title': task.title or '',
        'teacher': (course.teachers or '').strip(),
        'term': course.courseTerm or '',
        'college': FIXED_COLLEGE,
        'major': FIXED_MAJOR,
    }


_PLACEHOLDER_RE = re.compile(r'\{\{\s*([^{}]+?)\s*\}\}')
_COMBO_RE = re.compile(
    r'^(?P<lead>\s*)(?P<a>报告人|姓名)\s*(?P<a_colon>[：:])(?P<a_val>.*?)'
    r'(?P<b>学号)\s*(?P<b_colon>[：:])(?P<b_val>.*?)'
    r'(?P<c>班级)\s*(?P<c_colon>[：:])(?P<c_val>.*)$'  # 允许行首缩进；末组贪婪
)


def _replace_char_range(paragraph, start: int, end: int, new_text: str) -> bool:
    """Replace [start, end) chars; keep run formatting (underline etc.)."""
    if start < 0 or end < start:
        return False
    runs = paragraph.runs
    if not runs:
        return False

    span = end - start
    if span > 0 and len(new_text) < span:
        fill = new_text + (" " * (span - len(new_text)))
    elif span > 0 and len(new_text) >= span:
        # 值写满原空白时仍留一空格，避免「学号：xxx班级：」粘连
        fill = new_text + " "
    else:
        fill = new_text

    if start == end:
        pos = 0
        for run in runs:
            text = run.text or ""
            n = len(text)
            if pos <= start <= pos + n:
                off = start - pos
                run.text = text[:off] + fill + text[off:]
                return True
            pos += n
        runs[-1].text = (runs[-1].text or "") + fill
        return True

    pos = 0
    inserted = False
    for run in runs:
        text = run.text or ""
        out = []
        for ch in text:
            if start <= pos < end:
                if not inserted and pos == start:
                    out.append(fill)
                    inserted = True
            else:
                out.append(ch)
            pos += 1
        run.text = "".join(out)
    return inserted


def _paragraph_full_text(paragraph) -> str:
    return ''.join(r.text or '' for r in paragraph.runs) if paragraph.runs else (paragraph.text or '')


def _apply_placeholders_preserve(paragraph, values: dict) -> None:
    text = _paragraph_full_text(paragraph)
    matches = list(_PLACEHOLDER_RE.finditer(text))
    if not matches:
        return
    # 从后往前替换，避免偏移
    for m in reversed(matches):
        key = m.group(1).strip()
        if key in values and values[key]:
            _replace_char_range(paragraph, m.start(), m.end(), str(values[key]))


def _fill_combo_preserve(paragraph, values: dict) -> bool:
    text = _paragraph_full_text(paragraph)
    m = _COMBO_RE.match(text)
    if not m:
        return False

    def fill_group(group_name: str, value: str, match) -> None:
        if not value:
            return
        old = match.group(group_name) or ""
        if old.strip():
            return
        gs = match.start(group_name)
        ge = match.end(group_name)
        if gs == ge:
            colon_key = group_name.replace("_val", "_colon")
            gs = ge = match.end(colon_key)
        _replace_char_range(paragraph, gs, ge, value)

    # back to front + re-match after each fill
    for g, key in (("c_val", "班级"), ("b_val", "学号"), ("a_val", "姓名")):
        text = _paragraph_full_text(paragraph)
        m = _COMBO_RE.match(text)
        if not m:
            break
        fill_group(g, values.get(key) or "", m)
    return True


def _fill_label_value_preserve(paragraph, label: str, value: str, *, overwrite: bool) -> None:
    """填写「标签：后面的内容」；overwrite=True 时覆盖已有文字（保留下划线 run）。"""
    if not value or label in SKIP_COVER_LABELS:
        return
    text = _paragraph_full_text(paragraph)
    if label in ('报告人', '姓名', '学号', '班级') and ('学号' in text and '班级' in text and ('报告人' in text or '姓名' in text)):
        return
    # 整行「标签：值……」——值区含原文与尾部空格
    pat = re.compile(rf'({re.escape(label)}\s*[：:])(.*)$')
    m = pat.search(text)
    if not m:
        return
    rest = m.group(2) or ''
    if not overwrite and rest.strip():
        return
    # 已是目标值则不动（避免无意义重写）
    if rest.strip() == value.strip():
        return
    _replace_char_range(paragraph, m.start(2), m.end(2), value)


def autofill_paragraph(paragraph, values: dict) -> None:
    text0 = _paragraph_full_text(paragraph)
    # 实验时间 / 提交时间整行跳过
    for skip in SKIP_COVER_LABELS:
        if text0.strip().startswith(skip):
            return

    _apply_placeholders_preserve(paragraph, values)
    _fill_combo_preserve(paragraph, values)

    # 强制用系统值覆盖（标准封面）
    for label, key in (
        ('课程名称', '课程名称'),
        ('实验项目名称', '实验名称'),
        ('实验名称', '实验名称'),
        ('学院', '学院'),
        ('专业', '专业'),
        ('指导教师', '指导教师'),
        ('指导老师', '指导教师'),
    ):
        _fill_label_value_preserve(paragraph, label, values.get(key) or '', overwrite=True)

    # 单行报告人/学号/班级（非组合行）仅补空白
    for label, key in (
        ('报告人', '姓名'),
        ('姓名', '姓名'),
        ('学号', '学号'),
        ('班级', '班级'),
    ):
        _fill_label_value_preserve(paragraph, label, values.get(key) or '', overwrite=False)


def _iter_paragraphs(document):
    for p in document.paragraphs:
        yield p
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def autofill_docx_bytes(src_abs_path: str, values: dict) -> bytes:
    doc = Document(src_abs_path)
    for p in _iter_paragraphs(doc):
        autofill_paragraph(p, values)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def student_download_filename(task, profile, user) -> str:
    parts = [
        safe_filename(user.username or ''),
        safe_filename(profile.name or ''),
        safe_filename(task.title or '报告'),
        '报告模板.docx',
    ]
    name = '_'.join(p for p in parts if p)
    if not name.lower().endswith('.docx'):
        name += '.docx'
    return name
